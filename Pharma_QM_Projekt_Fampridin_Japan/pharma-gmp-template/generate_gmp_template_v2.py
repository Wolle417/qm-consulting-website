"""
Optimized GMP Template Generator v2
- Cleaner title page with less whitespace
- No Revision History page (comes from MD file)
- No TOC page (comes from MD file)
- Just: Title Page + Content Marker
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_margins(section, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """Set page margins in cm."""
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_page_border(paragraph, color="E0E0E0"):
    """Add a subtle bottom border to a paragraph (used for header)."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_styles(doc):
    """Define Normal, Heading 1/2/3 styles."""
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1E, 0x60, 0x91)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def add_header_footer(section):
    """Create GMP header and footer with page numbers."""
    # HEADER
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    run_left = para.add_run("Document No.: [DOC-ID]     ")
    run_left.font.size = Pt(9)
    run_left.font.name = 'Calibri'

    run_right = para.add_run("Version: [X.X] – [Date]")
    run_right.font.size = Pt(9)
    run_right.font.name = 'Calibri'

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_page_border(para)

    # FOOTER
    footer = section.footer
    f = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    r_left = f.add_run("Confidential – Internal Use Only    ")
    r_left.font.size = Pt(9)
    r_left.font.name = 'Calibri'

    r_right = f.add_run("Page ")
    r_right.font.size = Pt(9)
    r_right.font.name = 'Calibri'
    
    # Add page number field
    add_page_number(f)

    f.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_page_number(paragraph):
    """Add page number field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'


def add_title_page(doc):
    """
    Generate a clean, compact GMP title page.
    REDUCED WHITESPACE compared to v1.
    """
    # Document Number
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[DOCUMENT NUMBER]")
    r.font.size = Pt(14)
    r.font.bold = True

    # 1 spacer instead of 2
    doc.add_paragraph("")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("[DOCUMENT TITLE]")
    rt.font.size = Pt(20)
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(0x1E, 0x60, 0x91)

    # 1 spacer
    doc.add_paragraph("")

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("[SUBTITLE / PRODUCT / PROJECT]")
    rs.font.size = Pt(14)

    # 2 spacers instead of 3
    for _ in range(2):
        doc.add_paragraph("")

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        "Version: [X.X]\n"
        "Date: [DD.MM.YYYY]\n"
        "Author: [Name]"
    )
    rm.font.size = Pt(11)

    # 3 spacers bottom instead of 5
    for _ in range(3):
        doc.add_paragraph("")

    # Note
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = note.add_run("This document is for demonstration and training purposes only.")
    n.font.size = Pt(9)
    n.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_content_start(doc):
    """
    Add content marker on NEW page.
    The MD content will be inserted starting here.
    """
    doc.add_page_break()
    p = doc.add_paragraph("{{CONTENT_START}}")
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
    p.runs[0].font.bold = True


def main():
    """Generate optimized template v2."""
    doc = Document()

    # Base page setup
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    set_margins(section)
    configure_styles(doc)
    add_header_footer(section)

    # Build template pages
    # ONLY Title Page + Content Marker
    # NO Revision History page (comes from MD)
    # NO TOC page (comes from MD)
    add_title_page(doc)
    add_content_start(doc)

    # Save
    out = "GMP_Template_v2.docx"
    doc.save(out)
    print(f"✅ Optimized template generated: {out}")
    print(f"   - Compact title page (less whitespace)")
    print(f"   - No Revision History (will come from MD)")
    print(f"   - No TOC (will come from MD)")


if __name__ == "__main__":
    main()
