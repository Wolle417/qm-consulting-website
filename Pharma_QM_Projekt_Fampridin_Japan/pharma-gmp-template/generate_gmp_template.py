from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------
# 1) HELPER FUNCTIONS
# ---------------------------------------------------------

def set_margins(section, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """Set page margins in cm."""
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_page_border(paragraph, color="E0E0E0"):
    """Add a subtle bottom border to a paragraph (used for header)."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_styles(doc):
    """Define Normal, Heading 1/2/3 and a table style."""
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1E, 0x60, 0x91)  # GMP blue
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


def add_header_footer(section):
    """Create GMP header and footer."""
    # HEADER
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    run_left = para.add_run("Document No.: [DOC-ID]     ")
    run_left.font.size = Pt(9)

    run_right = para.add_run("Version: [X.X] – [Date]")
    run_right.font.size = Pt(9)

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_page_border(para)

    # FOOTER
    footer = section.footer
    f = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    r_left = f.add_run("Confidential – Internal Use Only    ")
    r_left.font.size = Pt(9)

    r_right = f.add_run("Page X of Y")
    r_right.font.size = Pt(9)

    f.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_title_page(doc):
    """Generate a clean GMP title page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[DOCUMENT NUMBER]")
    r.font.size = Pt(14)
    r.font.bold = True

    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("[DOCUMENT TITLE]")
    rt.font.size = Pt(20)
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(0x1E, 0x60, 0x91)

    doc.add_paragraph("")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("[SUBTITLE / PRODUCT / PROJECT]")
    rs.font.size = Pt(14)

    # Spacer
    for _ in range(3):
        doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        "Version: [X.X]\n"
        "Date: [DD.MM.YYYY]\n"
        "Author: [Name]"
    )
    rm.font.size = Pt(11)

    # Spacer bottom
    for _ in range(5):
        doc.add_paragraph("")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = note.add_run("This document is for demonstration and training purposes only.")
    n.font.size = Pt(9)
    n.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_revision_page(doc):
    """Add a Revision History page with placeholder."""
    doc.add_page_break()
    h = doc.add_paragraph("Revision History")
    h.style = "Heading 1"

    t = doc.add_table(rows=1, cols=4)
    hdr = t.rows[0].cells
    hdr[0].text = "Version"
    hdr[1].text = "Date"
    hdr[2].text = "Changes"
    hdr[3].text = "Author"


def add_toc_page(doc):
    doc.add_page_break()
    h = doc.add_paragraph("Table of Contents")
    h.style = "Heading 1"
    doc.add_paragraph(
        "Insert automatic TOC here (Word > References > Table of Contents)."
    )


def add_content_start(doc):
    doc.add_page_break()
    p = doc.add_paragraph("{{CONTENT_START}}")
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
    p.runs[0].font.bold = True


# ---------------------------------------------------------
# MAIN
# ---------------------------------------------------------

def main():
    doc = Document()

    # Base page setup
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    set_margins(section)
    configure_styles(doc)
    add_header_footer(section)

    # Build template pages
    add_title_page(doc)
    add_revision_page(doc)
    add_toc_page(doc)
    add_content_start(doc)

    # Save
    out = "GMP_Template_v1.docx"
    doc.save(out)
    print(f"Template generated: {out}")


if __name__ == "__main__":
    main()
