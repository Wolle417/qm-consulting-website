"""
MD to Word Converter with GMP Template
Uses the GMP_Template_v1.docx as base template.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------
# HELPER FUNCTIONS (from generate_gmp_template.py)
# ---------------------------------------------------------

def set_margins(section, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """Set page margins in cm."""
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_page_border(paragraph, color="E0E0E0"):
    """Add a subtle bottom border to a paragraph (used for header)."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_styles(doc):
    """Define Normal, Heading 1/2/3 styles."""
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1E, 0x60, 0x91)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def add_header_footer(section, doc_number, version, date):
    """Update GMP header and footer with actual document info."""
    # HEADER
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()

    run_left = para.add_run(f"Document No.: {doc_number}     ")
    run_left.font.size = Pt(9)
    run_left.font.name = 'Calibri'

    run_right = para.add_run(f"Version: {version} – {date}")
    run_right.font.size = Pt(9)
    run_right.font.name = 'Calibri'

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_page_border(para)

    # FOOTER
    footer = section.footer
    f = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    f.clear()

    r_left = f.add_run("Confidential – Internal Use Only    ")
    r_left.font.size = Pt(9)
    r_left.font.name = 'Calibri'

    r_right = f.add_run("Page ")
    r_right.font.size = Pt(9)
    r_right.font.name = 'Calibri'
    
    # Add page number field
    add_page_number(f)

    f.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_page_number(paragraph):
    """Add actual page number field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'


def parse_metadata(content):
    """Extract document metadata from Markdown."""
    metadata = {
        'doc_number': 'DOC-001',
        'version': 'v1.0',
        'date': '2025',
        'title': 'Document',
        'subtitle': '',
    }
    
    lines = content.split('\n')
    for i, line in enumerate(lines[:30]):
        if 'Document Number:' in line or '**Document Number:**' in line:
            metadata['doc_number'] = line.split(':')[-1].strip().replace('*', '')
        elif 'Version:' in line or '**Version:**' in line:
            metadata['version'] = line.split(':')[-1].strip().replace('*', '')
        elif 'Effective Date:' in line or 'Date:' in line or '**Effective Date:**' in line:
            date_str = line.split(':')[-1].strip().replace('*', '')
            date_str = date_str.replace('2026', '2025')  # Fix year
            metadata['date'] = date_str
        elif line.startswith('# ') and not metadata.get('title_set'):
            metadata['title'] = line[2:].strip()
            metadata['title_set'] = True
        elif line.startswith('## ') and not metadata.get('subtitle_set') and i < 5:
            metadata['subtitle'] = line[3:].strip()
            metadata['subtitle_set'] = True
    
    return metadata


def replace_title_page(doc, metadata):
    """Replace placeholders in title page."""
    for para in doc.paragraphs[:20]:  # Only check first 20 paragraphs (title page area)
        for run in para.runs:
            if '[DOCUMENT NUMBER]' in run.text:
                run.text = run.text.replace('[DOCUMENT NUMBER]', metadata['doc_number'])
            if '[DOCUMENT TITLE]' in run.text:
                run.text = run.text.replace('[DOCUMENT TITLE]', metadata['title'])
            if '[SUBTITLE / PRODUCT / PROJECT]' in run.text:
                subtitle = metadata['subtitle'] if metadata['subtitle'] else f"Pharmaceutical Quality Documentation\nFampridin (4-Aminopyridine)"
                run.text = run.text.replace('[SUBTITLE / PRODUCT / PROJECT]', subtitle)
            if 'Version: [X.X]' in run.text:
                run.text = run.text.replace('Version: [X.X]', f"Version: {metadata['version']}")
            if 'Date: [DD.MM.YYYY]' in run.text:
                run.text = run.text.replace('Date: [DD.MM.YYYY]', f"Date: {metadata['date']}")
            if 'Author: [Name]' in run.text:
                run.text = run.text.replace('Author: [Name]', "Author: Stefan Schönwälder")


def find_content_start(doc):
    """Find the {{CONTENT_START}} marker paragraph."""
    for i, para in enumerate(doc.paragraphs):
        if '{{CONTENT_START}}' in para.text:
            return i
    return -1


def convert_md_to_docx(md_file, output_file, template_path):
    """Convert a Markdown file to Word using GMP template."""
    print(f"Converting: {os.path.basename(md_file)}")
    
    # Read MD file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Parse metadata
    metadata = parse_metadata(md_content)
    
    # Load template
    doc = Document(template_path)
    
    # Configure styles (ensure consistency)
    configure_styles(doc)
    
    # Update header/footer with actual document info
    section = doc.sections[0]
    add_header_footer(section, metadata['doc_number'], metadata['version'], metadata['date'])
    
    # Replace title page placeholders
    replace_title_page(doc, metadata)
    
    # Find where to insert content
    content_start_idx = find_content_start(doc)
    if content_start_idx == -1:
        print(f"  ⚠ Warning: {{{{CONTENT_START}}}} marker not found, appending at end")
        content_start_idx = len(doc.paragraphs)
    else:
        # Remove the marker paragraph
        p = doc.paragraphs[content_start_idx]
        p._element.getparent().remove(p._element)
    
    # Process markdown content
    lines = md_content.split('\n')
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []
    skip_first_h1 = True  # Skip the first # heading (it's the title, already in template)
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip frontmatter-like lines and document header
        if i < 50 and ('**Document Number:**' in line or '**Version:**' in line or 
                       '**Effective Date:**' in line or '---' == line.strip() or
                       '**Prepared by:**' in line or '**Reviewed by:**' in line or
                       '**Approved by:**' in line or '**Product:**' in line or
                       '**Status:**' in line or '**Project:**' in line):
            i += 1
            continue
        
        # Skip the first main heading (title) as it's already in the template
        if skip_first_h1 and line.startswith('# ') and not line.startswith('## '):
            skip_first_h1 = False
            i += 1
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:].strip())
            p.style = 'Heading 1'
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:].strip())
            p.style = 'Heading 2'
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:].strip())
            p.style = 'Heading 3'
        
        # Tables
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
        elif in_table and not ('|' in line):
            create_table_from_md(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Regular paragraph
        elif line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        else:
            if not in_table:
                doc.add_paragraph()
        
        i += 1
    
    # Handle remaining table
    if table_rows:
        create_table_from_md(doc, table_rows)
    
    # Save
    doc.save(output_file)
    print(f"  ✓ Created: {os.path.basename(output_file)}")


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code)."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def create_table_from_md(doc, table_rows):
    """Create a Word table from markdown table rows."""
    if len(table_rows) < 2:
        return
    
    rows = []
    for row in table_rows:
        if '---' in row or '===' in row:
            continue
        cells = [cell.strip() for cell in row.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                if i == 0:  # Header row
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
    
    doc.add_paragraph()


def main():
    """Convert all 7 main documents using the template."""
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(base_dir)
    template_path = os.path.join(base_dir, "GMP_Template_v1.docx")
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        print("Please run generate_gmp_template.py first!")
        return
    
    documents = [
        ("01_Quality_Manual/QM-001_Quality_Manual_v1.1.md", 
         "01_Quality_Manual/QM-001_Quality_Manual_v1.1.docx"),
        ("02_Master_Validation_Plan/MVP-001_Master_Validation_Plan_v1.0.md",
         "02_Master_Validation_Plan/MVP-001_Master_Validation_Plan_v1.0.docx"),
        ("03_Regulatory_Process_Overview/REG-PROC-001_Regulatory_Process_API_Export_Japan_Fampridin_v1.0.md",
         "03_Regulatory_Process_Overview/REG-PROC-001_Regulatory_Process_API_Export_Japan_Fampridin_v1.0.docx"),
        ("04_Process_Validation/PVP-FAM-001_Process_Validation_Protocol_Fampridin_v1.0.md",
         "04_Process_Validation/PVP-FAM-001_Process_Validation_Protocol_Fampridin_v1.0.docx"),
        ("05_Analytical_Method_Validation/AMV-HPLC-001_Analytical_Method_Validation_HPLC_Fampridin_v1.0.md",
         "05_Analytical_Method_Validation/AMV-HPLC-001_Analytical_Method_Validation_HPLC_Fampridin_v1.0.docx"),
        ("06_CAPA_Report/CAPA-2025-001_HVAC_Overpressure_Deviation_v1.0.md",
         "06_CAPA_Report/CAPA-2025-001_HVAC_Overpressure_Deviation_v1.0.docx"),
        ("07_DMF_Drug_Master_File/DMF-FAM-001_Module_3.2.S_Fampridin_v1.0.md",
         "07_DMF_Drug_Master_File/DMF-FAM-001_Module_3.2.S_Fampridin_v1.0.docx"),
    ]
    
    print("="*60)
    print("MD to Word Converter - Using GMP Template v1")
    print("="*60)
    print()
    
    success_count = 0
    for md_path, docx_path in documents:
        try:
            md_full_path = os.path.join(project_dir, md_path)
            docx_full_path = os.path.join(project_dir, docx_path)
            
            if os.path.exists(md_full_path):
                convert_md_to_docx(md_full_path, docx_full_path, template_path)
                success_count += 1
            else:
                print(f"✗ Not found: {md_path}")
        except Exception as e:
            print(f"✗ Error converting {md_path}: {str(e)}")
            import traceback
            traceback.print_exc()
    
    print()
    print("="*60)
    print(f"Conversion complete: {success_count}/{len(documents)} documents")
    print(f"Template used: GMP_Template_v1.docx")
    print(f"Output: Word files in their respective source folders")
    print("="*60)


if __name__ == "__main__":
    main()
